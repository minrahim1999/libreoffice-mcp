"""Document (docx) operations for LibreOffice MCP."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def create_document(output_path: str | Path) -> None:
    doc = Document()
    doc.save(str(output_path))


def add_heading(file_path: str | Path, text: str, *, level: int = 1) -> None:
    doc = Document(str(file_path))
    doc.add_heading(text, level=level)
    doc.save(str(file_path))


def add_paragraph(
    file_path: str | Path,
    text: str,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
    font_size: int | None = None,
) -> None:
    doc = Document(str(file_path))
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    doc.save(str(file_path))


def add_document_table(
    file_path: str | Path,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    doc = Document(str(file_path))
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for j, header in enumerate(headers):
        hdr_cells[j].text = header
        for paragraph in hdr_cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for j, val in enumerate(row_data):
            row_cells[j].text = val

    doc.save(str(file_path))
